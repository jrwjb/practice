import re
import os
import copy
import pandas as pd
# import docx
from docx import Document
from docx.shared import Cm
from docx.shared import RGBColor
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from hyperlink import add_hyperlink
from record_doc import record_doc


target_path = "E:/Project/Protein/Label_free/P20190600772"
os.chdir(target_path)


document = Document('C:\\Users\\jbwang\\Desktop\\Label free.docx')
paragraphs = document.paragraphs
# styles = document.styles
# groupvs_list = record_doc()[10].split('、')

def paragraph_header(pa, size, family, r = 0x00, g = 0x00, b = 0x00, bold = None):
    pa.font.size = Pt(size)
    pa.font.name = family
    if bold == True:
        pa.font.bold = True
    pa.font.color.rgb = RGBColor(r, g, b)
    pa._element.rPr.rFonts.set(qn('w:eastAsia'), family)

## 报告首页
for i in range(4):
    pa = paragraphs[i+25].add_run(str(record_doc()[i]))
    paragraph_header(pa, size=14, family=u'微软雅黑', bold=True)

## 蛋白鉴定统计
pro_df = pd.read_excel('E:\\Project\\Protein\\Label_free\\P20190600772\\报告及附件\\质谱鉴定和定量结果\\附件1_蛋白质鉴定列表.xlsx', sheet_name=0)
pro_total_num = len(pro_df.Protein)
# sample_list = [i for i in pro_df.columns if re.search(r'iBAQ|intensity', i, re.I)]
summary = {i.split(' ')[-1]: pro_df[i].notnull().sum() for i in pro_df.columns if re.search(r'iBAQ|intensity', i, re.I)}
summary_df = pd.DataFrame.from_dict(summary, orient='index').reset_index()
summary_df.loc[summary_df.shape[0]+1] = [f'{summary_df.shape[0]}组共有', pro_total_num]
summary_df.insert(0, 'database', record_doc()[5])

## 插入表格
tables = document.tables
for row_num in range(summary_df.shape[0]):
    row_line = tables[0].add_row()
    row_line.height = Cm(0.7)
    row_cells = row_line.cells
    for col_num in range(summary_df.shape[1]):
        tmp = summary_df.iloc[row_num, col_num]
        # if type(tmp) != str and tmp.dtype == 'float64':
        #     tmp = round(tmp, 3)
        pa = row_cells[col_num].paragraphs[0].add_run(str(tmp))
        row_cells[col_num].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[col_num].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_header(pa, size=10.5, family="Times New Roman")      #### family = 'Calibri'

tmp_text = tables[1].cell(8, 1).text.replace('xxx', record_doc()[6])
tables[1].cell(8, 1).text = ''
p = tables[1].cell(8, 1).paragraphs[0].add_run(tmp_text)
tables[1].cell(8, 1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
tables[1].cell(8, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
paragraph_header(p, size=10.5, family="Calibri")

# print()
def insert_png(png1, png2):
    p.text = p.text.strip().replace(png1, "")
    run = p.add_run()
    run.add_picture(png2, width=Inches(4.5))


for p in paragraphs:
    if p.text.startswith('本项目'):
        p.text = re.sub('xxx', record_doc()[4], p.text, 1).replace('xxx', str(pro_total_num))

    if "[venn]" in p.text:
        insert_png("[venn]", "报告及附件/Evaluation/Venn/组内/venn_DMSO-14-3-3OE.png")

    if "[venn_group]" in p.text:
        insert_png("[venn_group]", "报告及附件/Evaluation/Venn/组间/4-Set-Venn.png")

    # if "[sample_info]" in p.text:
    #     p.text = p.text.strip().replace("[sample_info]", "")
    #     run = p.add_run(record_doc()[7])

    if "[mass_error]" in p.text:
        insert_png("[mass_error]", "报告及附件/Evaluation/mass_error_distribution.png")

    if "[pep_score]" in p.text:
        insert_png("[pep_score]", "报告及附件/Evaluation/Andromeda_Score_Distribution.png")

    if "[protein_MM]" in p.text:
        insert_png("[protein_MM]", "报告及附件/Evaluation/Molecular_Weight_Distribution.png")

    if "[pep_len]" in p.text:
        insert_png("[pep_len]", "报告及附件/Evaluation/Peptide_Length_Distribution.png")

    if "[pep_coverage]" in p.text:
        insert_png("[pep_coverage]", "报告及附件/Evaluation/Protein_Sequence_Coverage_Distribution.png")

    if "[pep_count]" in p.text:
        insert_png("[pep_count]", "报告及附件/Evaluation/Peptide_Count_Distribution.png")


document.save('C:\\Users\\jbwang\\Desktop\\Label free相对定量蛋白质组学生物信息学分析报告.docx')


